"""Inspect the new CV docx + pdf so we know exactly what the HTML has to reproduce."""
import os

import docx

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX = os.path.join(REPO, "assets", "Cory Boris Curriculum Vitae.docx")
PDF = os.path.join(REPO, "assets", "Cory Boris Curriculum Vitae.pdf")

d = docx.Document(DOCX)

print("=== SECTION / PAGE SETUP ===")
for s in d.sections:
    print("  page   %.4f x %.4f in" % (s.page_width.inches, s.page_height.inches))
    print("  margins L%.4f R%.4f T%.4f B%.4f in"
          % (s.left_margin.inches, s.right_margin.inches,
             s.top_margin.inches, s.bottom_margin.inches))

print("\n=== DEFAULT STYLE ===")
normal = d.styles["Normal"]
print("  font: %s  size: %s" % (normal.font.name, normal.font.size.pt if normal.font.size else None))

print("\n=== BODY ELEMENTS (in document order) ===")
body = d.element.body
ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
kinds = []
for child in body:
    tag = child.tag.replace(ns, "")
    kinds.append(tag)
print("  element sequence:", " ".join(kinds))

print("\n=== PARAGRAPHS ===")
fonts = set()
sizes = set()
for i, p in enumerate(d.paragraphs):
    text = p.text
    pf = p.paragraph_format
    bits = []
    if pf.left_indent is not None:
        bits.append("L=%.4fin" % pf.left_indent.inches)
    if pf.first_line_indent is not None:
        bits.append("ind=%.4fin" % pf.first_line_indent.inches)
    if p.alignment is not None:
        bits.append("align=%s" % p.alignment)
    if pf.line_spacing is not None:
        bits.append("ls=%s" % pf.line_spacing)
    # tab stops drive the right-aligned dates
    tabs = []
    for t in pf.tab_stops:
        tabs.append("%s@%.3fin" % (t.alignment, t.position.inches))
    if tabs:
        bits.append("tabs[%s]" % ",".join(tabs))
    for r in p.runs:
        if r.font.name:
            fonts.add(r.font.name)
        if r.font.size:
            sizes.add(r.font.size.pt)
    has_tab = "\t" in "".join(r.text for r in p.runs)
    print("%3d %-4s %-58s %s" % (i, "TAB" if has_tab else "", (text[:56] or "<empty>"), " ".join(bits)))

print("\nfonts used: %s" % sorted(fonts))
print("sizes used: %s" % sorted(sizes))
print("paragraph count: %d, table count: %d" % (len(d.paragraphs), len(d.tables)))

print("\n=== PDF ===")
try:
    import fitz
    doc = fitz.open(PDF)
    print("  pages: %d" % doc.page_count)
    for pno in range(doc.page_count):
        page = doc[pno]
        print("  page %d: %.2f x %.2f pt (%.3f x %.3f in)"
              % (pno + 1, page.rect.width, page.rect.height,
                 page.rect.width / 72.0, page.rect.height / 72.0))
    pdf_fonts = set()
    for pno in range(doc.page_count):
        for f in doc[pno].get_fonts():
            pdf_fonts.add(f[3])
    print("  embedded fonts: %s" % sorted(pdf_fonts))
except Exception as e:
    print("  pdf inspect failed:", e)
